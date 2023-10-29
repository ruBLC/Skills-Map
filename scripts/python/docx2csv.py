from docx import Document
import csv
import os

# Функция для извлечения всех таблиц из DOCX файла и сохранения их в CSV
def docx_to_csv_in_folder(docx_filename):
    # Шаг 1: Чтение файла DOCX
    doc = Document(docx_filename)
    table_count = 0
    folder_name = None
    
    # Поиск строки после "ПРОФЕССИОНАЛЬНЫЙ СТАНДАРТ"
    for para in doc.paragraphs:
        if folder_name is not None:
            folder_name = para.text
            break
        if "ПРОФЕССИОНАЛЬНЫЙ СТАНДАРТ" in para.text:
            folder_name = True
    
    # Если нашли строку для названия папки, создаем папку
    if folder_name and folder_name is not True:
        if not os.path.exists(folder_name):
            os.makedirs(folder_name)
    
    for table in doc.tables:
        # Шаг 2: Извлечение таблицыc
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            data.append(row_data)
        
        # Шаг 3: Сохранение таблицы в CSV в формате UTF-8
        csv_filename = f"{folder_name}/table_{table_count}.csv" if folder_name else f"table_{table_count}.csv"
        with open(csv_filename, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f, quoting=csv.QUOTE_ALL)  # Все поля будут заключены в кавычки
            writer.writerows(data)
        
        print(f"Таблица сохранена в {csv_filename} в формате UTF-8 с кавычками")
        table_count += 1

# Пример использования функции
docx_to_csv_in_folder(r"C:\Users\mitin\Downloads\file_108240.docx")

